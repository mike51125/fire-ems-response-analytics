# -*- coding: utf-8 -*-
"""
Created on Tue May  5 12:13:08 2026

@author: mike5
"""


import pandas as pd
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# =====================================================
# FILE PATHS
# =====================================================
call_log_file = r"C:\Users\mike5\Desktop\Board Reports\April\April CallLog.csv"
monthly_incident_file = r"C:\Users\mike5\Desktop\Board Reports\April\MonthlyIncidentNumbersApril.csv"
overlap_file = r"C:\Users\mike5\Desktop\Board Reports\April\Overlapping Incidents April 2026.xlsx"

chart_path = r"C:\Users\mike5\Desktop\Board Reports\April\calls_by_hour.png"
output_doc_path = r"C:\Users\mike5\Desktop\Board Reports\April\D7_Response_Analysis.docx"

# =====================================================
# LOAD CALL LOG AND MONTHLY INCIDENT LIST
# =====================================================
df = pd.read_csv(call_log_file)
monthly_df = pd.read_csv(monthly_incident_file)

df.columns = df.columns.str.strip()
monthly_df.columns = monthly_df.columns.str.strip()

# =====================================================
# COLUMN NAMES
# =====================================================
incident_col = "Core incident number"
station_col = "Station"
aid_col = "Core aid direction"
response_col = "Unit response time"
response_mode_col = "Unit response mode to scene"
date_col = "Date"
time_col = "Time"

# =====================================================
# CLEAN INCIDENT NUMBERS AND STATIONS
# =====================================================
df[incident_col] = df[incident_col].astype(str).str.strip()
monthly_df[incident_col] = monthly_df[incident_col].astype(str).str.strip()

df[station_col] = df[station_col].astype(str).str.strip()
monthly_df[station_col] = monthly_df[station_col].astype(str).str.strip()

# =====================================================
# REMOVE DUPLICATE INCIDENT NUMBERS FROM CALL LOG
# This cleaned call log is used for all performance metrics.
# =====================================================
raw_record_count = len(df)
duplicate_count = df[incident_col].duplicated().sum()

df = df.drop_duplicates(subset=[incident_col], keep="first").copy()

deduped_record_count = len(df)

# =====================================================
# REMOVE DUPLICATE INCIDENT NUMBERS FROM MONTHLY LIST
# Monthly list is the source of truth for total call volume.
# =====================================================
monthly_raw_record_count = len(monthly_df)
monthly_duplicate_count = monthly_df[incident_col].duplicated().sum()

monthly_df = monthly_df.drop_duplicates(subset=[incident_col], keep="first").copy()

monthly_deduped_record_count = len(monthly_df)

# =====================================================
# IDENTIFY CALLS ON MONTHLY LIST BUT NOT IN CALL LOG
# These calls count toward station workload, but are excluded
# from response time, overlap, duration, and performance metrics.
# =====================================================
missing_calls = monthly_df[
    ~monthly_df[incident_col].isin(df[incident_col])
].copy()

missing_calls["Data Status"] = "Insufficient Data - Excluded From Performance Metrics"

missing_data_call_count = missing_calls[incident_col].nunique()

missing_by_station = (
    missing_calls.groupby(station_col)[incident_col]
    .nunique()
    .reset_index()
)

missing_by_station.columns = ["Station", "Insufficient Data Calls"]

# =====================================================
# CLEAN CALL LOG DATA FOR PERFORMANCE METRICS
# =====================================================
df[aid_col] = df[aid_col].astype(str).str.strip().str.lower()
df[response_mode_col] = df[response_mode_col].astype(str).str.strip().str.lower()

df[response_col] = (
    df[response_col]
    .astype(str)
    .str.replace(",", "", regex=False)
)

df[response_col] = pd.to_numeric(df[response_col], errors="coerce")

df["Call DateTime"] = pd.to_datetime(
    df[date_col].astype(str) + " " + df[time_col].astype(str),
    errors="coerce"
)

df["Hour"] = df["Call DateTime"].dt.hour

# =====================================================
# FORMAT SECONDS AS MM:SS
# =====================================================
def seconds_to_mmss(seconds):
    if pd.isna(seconds):
        return ""
    seconds = int(round(seconds))
    minutes = seconds // 60
    secs = seconds % 60
    return f"{minutes}:{secs:02d}"

# =====================================================
# SAFE PERCENT HELPER
# =====================================================
def safe_percent(numerator, denominator):
    if denominator == 0 or pd.isna(denominator):
        return 0
    return (numerator / denominator) * 100

# =====================================================
# WORD TABLE HELPER
# =====================================================
def add_table_from_df(document, dataframe, title=None):
    if title:
        document.add_heading(title, level=1)

    table = document.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(dataframe.columns):
        hdr_cells[i].text = str(col)

    for _, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    return table

# =====================================================
# CALL TYPE FILTERS
# Call log only; missing monthly calls are excluded here.
# =====================================================
first_due_df = df[
    (df[aid_col] == "received") |
    (df[aid_col] == "na") |
    (df[aid_col] == "nan") |
    (df[aid_col] == "")
].copy()

mutual_aid_df = df[df[aid_col] == "given"].copy()

first_due_emergent_df = first_due_df[
    first_due_df[response_mode_col] == "emergent"
].copy()

# =====================================================
# MONTHLY CALL COUNTS
# Total monthly calls uses monthly incident list.
# Other metrics use call log only.
# =====================================================
total_monthly_calls = monthly_df[incident_col].nunique()
total_call_log_calls = df[incident_col].nunique()

total_first_due_calls = first_due_df[incident_col].nunique()
total_first_due_emergent_calls = first_due_emergent_df[incident_col].nunique()

# =====================================================
# TOTAL CALLS BY STATION
# Uses monthly list so missing/canceled/no-time calls are credited.
# =====================================================
total_calls_by_station = (
    monthly_df.groupby(station_col)[incident_col]
    .nunique()
    .reset_index()
)

total_calls_by_station.columns = ["Station", "Total Calls"]

# =====================================================
# TOTAL 1ST DUE CALLS BY STATION
# Call log only; missing calls are excluded from this metric.
# =====================================================
first_due_calls_by_station = (
    first_due_df.groupby(station_col)[incident_col]
    .nunique()
    .reset_index()
)

first_due_calls_by_station.columns = ["Station", "1st Due Calls"]

# =====================================================
# TOTAL 1ST DUE EMERGENT CALLS BY STATION
# Call log only; missing calls are excluded from this metric.
# =====================================================
first_due_emergent_calls_by_station = (
    first_due_emergent_df.groupby(station_col)[incident_col]
    .nunique()
    .reset_index()
)

first_due_emergent_calls_by_station.columns = [
    "Station",
    "1st Due Emergent Calls"
]

# =====================================================
# MUTUAL AID CALLS
# Call log only.
# =====================================================
total_mutual_aid_calls = mutual_aid_df[incident_col].nunique()

mutual_aid_calls_by_station = (
    mutual_aid_df.groupby(station_col)[incident_col]
    .nunique()
    .reset_index()
)

mutual_aid_calls_by_station.columns = ["Station", "Mutual Aid Calls"]

# =====================================================
# AVERAGE + 90TH PERCENTILE 1ST DUE EMERGENT RESPONSE TIME
# Call log only.
# =====================================================
avg_first_due_response_time = seconds_to_mmss(
    first_due_emergent_df[response_col].mean()
)

p90_first_due_response_time = seconds_to_mmss(
    first_due_emergent_df[response_col].quantile(0.90)
)

response_time_by_station = (
    first_due_emergent_df.groupby(station_col)[response_col]
    .agg(
        Average_Response_Time_Seconds="mean",
        P90_Response_Time_Seconds=lambda x: x.quantile(0.90)
    )
    .reset_index()
)

response_time_by_station["Average 1st Due Emergent Response Time"] = (
    response_time_by_station["Average_Response_Time_Seconds"]
    .apply(seconds_to_mmss)
)

response_time_by_station["90th Percentile 1st Due Emergent Response Time"] = (
    response_time_by_station["P90_Response_Time_Seconds"]
    .apply(seconds_to_mmss)
)

response_time_by_station_clean = response_time_by_station[
    [
        station_col,
        "Average 1st Due Emergent Response Time",
        "90th Percentile 1st Due Emergent Response Time"
    ]
].copy()

response_time_by_station_clean.columns = [
    "Station",
    "Average 1st Due Emergent Response Time",
    "90th Percentile 1st Due Emergent Response Time"
]

# =====================================================
# 1ST DUE EMERGENT RESPONSE TIME > 480
# Call log only.
# =====================================================
first_due_over_480 = first_due_emergent_df[
    first_due_emergent_df[response_col] > 480
].copy()

count_first_due_over_480 = first_due_over_480[incident_col].nunique()

percent_first_due_over_480 = safe_percent(
    count_first_due_over_480,
    total_first_due_emergent_calls
)

count_first_due_over_480_by_station = (
    first_due_over_480.groupby(station_col)[incident_col]
    .nunique()
    .reset_index()
)

count_first_due_over_480_by_station.columns = [
    "Station",
    "1st Due Emergent Calls >480"
]

percent_over_480_by_station = first_due_emergent_calls_by_station.merge(
    count_first_due_over_480_by_station,
    on="Station",
    how="left"
)

percent_over_480_by_station["1st Due Emergent Calls >480"] = (
    percent_over_480_by_station["1st Due Emergent Calls >480"]
    .fillna(0)
    .astype(int)
)

percent_over_480_by_station["Percent >480 Numeric"] = (
    percent_over_480_by_station["1st Due Emergent Calls >480"] /
    percent_over_480_by_station["1st Due Emergent Calls"]
) * 100

percent_over_480_by_station["Percent >480 Numeric"] = (
    percent_over_480_by_station["Percent >480 Numeric"]
    .fillna(0)
)

percent_over_480_by_station["Percent >480"] = (
    percent_over_480_by_station["Percent >480 Numeric"]
    .round(1)
    .astype(str) + "%"
)

percent_over_480_by_station = percent_over_480_by_station.drop(
    columns=["Percent >480 Numeric"]
)

# =====================================================
# DATAFRAME OF 1ST DUE EMERGENT RESPONSE TIME >480
# =====================================================
first_due_over_480_detail = first_due_over_480[
    [
        incident_col,
        station_col,
        response_mode_col,
        response_col
    ]
].copy()

first_due_over_480_detail.columns = [
    "Incident Number",
    "Station",
    "Response Mode",
    "Response Time Seconds"
]

first_due_over_480_detail["Incident Number"] = (
    first_due_over_480_detail["Incident Number"]
    .astype(str)
    .str.strip()
)

first_due_over_480_detail["Response Time mm:ss"] = (
    first_due_over_480_detail["Response Time Seconds"]
    .apply(seconds_to_mmss)
)

# =====================================================
# READ OVERLAPPING CALL FILE
# =====================================================
overlap_df = pd.read_excel(
    overlap_file,
    sheet_name="Fire Incidents",
    header=8
)

overlap_df.columns = overlap_df.columns.str.strip()

overlap_df["Incident Number"] = (
    overlap_df["Incident Number"]
    .astype(str)
    .str.strip()
)

overlap_df["Overlapping"] = pd.to_numeric(
    overlap_df["Overlapping"],
    errors="coerce"
).fillna(0)

overlap_df = overlap_df[
    [
        "Incident Number",
        "Overlapping"
    ]
].copy()

overlap_df = overlap_df.drop_duplicates(subset=["Incident Number"], keep="first")

# =====================================================
# COMPARE >480 FIRST DUE EMERGENT CALLS TO OVERLAPPING CALL FILE
# =====================================================
first_due_over_480_with_overlap = first_due_over_480_detail.merge(
    overlap_df,
    on="Incident Number",
    how="left"
)

first_due_over_480_with_overlap["Overlapping"] = (
    first_due_over_480_with_overlap["Overlapping"]
    .fillna(0)
)

first_due_over_480_with_overlap["Has Overlap"] = (
    first_due_over_480_with_overlap["Overlapping"] > 0
)

first_due_over_480_count = len(first_due_over_480_with_overlap)

first_due_over_480_with_any_overlap = (
    first_due_over_480_with_overlap["Has Overlap"]
    .sum()
)

percent_first_due_over_480_with_overlap = safe_percent(
    first_due_over_480_with_any_overlap,
    first_due_over_480_count
)

average_overlap_count_for_first_due_over_480 = (
    first_due_over_480_with_overlap["Overlapping"]
    .mean()
)

first_due_over_480_overlap_by_station = (
    first_due_over_480_with_overlap
    .groupby("Station")
    .agg(
        First_Due_Emergent_Calls_Over_480=("Incident Number", "count"),
        Calls_With_Overlap=("Has Overlap", "sum"),
        Avg_Overlapping_Calls=("Overlapping", "mean")
    )
    .reset_index()
)

first_due_over_480_overlap_by_station["Percent_With_Overlap"] = (
    first_due_over_480_overlap_by_station["Calls_With_Overlap"] /
    first_due_over_480_overlap_by_station["First_Due_Emergent_Calls_Over_480"]
) * 100

first_due_over_480_overlap_by_station_clean = (
    first_due_over_480_overlap_by_station.copy()
)

first_due_over_480_overlap_by_station_clean.columns = [
    "Station",
    "1st Due Emergent Calls >480",
    "Calls w/ Overlap",
    "Avg Overlapping Calls",
    "% w/ Overlap"
]

first_due_over_480_overlap_by_station_clean["Avg Overlapping Calls"] = (
    first_due_over_480_overlap_by_station_clean["Avg Overlapping Calls"]
    .round(2)
)

first_due_over_480_overlap_by_station_clean["% w/ Overlap"] = (
    first_due_over_480_overlap_by_station_clean["% w/ Overlap"]
    .round(1)
    .astype(str) + "%"
)

# =====================================================
# SYSTEM STRESS HOURS
# System stress = overlapping calls >= 2
# Used for chart shading only; no table in Word doc.
# Call log only.
# =====================================================
df_for_stress = df[[incident_col, "Hour"]].copy()

df_for_stress.columns = ["Incident Number", "Hour"]

df_for_stress["Incident Number"] = (
    df_for_stress["Incident Number"]
    .astype(str)
    .str.strip()
)

df_for_stress = df_for_stress.merge(
    overlap_df,
    on="Incident Number",
    how="left"
)

df_for_stress["Overlapping"] = df_for_stress["Overlapping"].fillna(0)

stress_hours = sorted(
    df_for_stress[
        (df_for_stress["Overlapping"] >= 2) &
        (df_for_stress["Hour"].notna())
    ]["Hour"].astype(int).unique()
)

stress_hours_df = pd.DataFrame({"System Stress Hours": stress_hours})

# =====================================================
# CALLS BY HOUR - CALL LOG ONLY
# Missing monthly calls do not have usable time data.
# =====================================================
calls_by_hour = (
    df.dropna(subset=["Hour"])
    .groupby("Hour")[incident_col]
    .nunique()
    .reindex(range(24), fill_value=0)
)

calls_by_hour.index = calls_by_hour.index.astype(int)

calls_by_hour_df = calls_by_hour.reset_index()
calls_by_hour_df.columns = ["Hour", "Call Count"]

# =====================================================
# CALLS BY HOUR COLUMN CHART WITH SYSTEM STRESS SHADING
# =====================================================
plt.figure(figsize=(8, 4.5))
ax = plt.gca()

for hour in stress_hours:
    ax.axvspan(hour - 0.5, hour + 0.5, alpha=0.15)

plt.bar(
    calls_by_hour.index,
    calls_by_hour.values,
    color="#253555",
    edgecolor="black"
)

plt.title("All Calls by Hour of Day")
plt.xlabel("Hour of Day")
plt.ylabel("Number of Calls")

even_hours = [h for h in range(24) if h % 2 == 0]
plt.xticks(even_hours)

plt.grid(False)
plt.gca().yaxis.grid(False)

plt.tight_layout()
plt.savefig(chart_path, dpi=300)
plt.show()
plt.close()

# =====================================================
# STATION SUMMARY TABLE FOR WORD REPORT
# Total Calls uses monthly list.
# Insufficient Data Calls are credited by station but excluded
# from all response-time and overlap metrics.
# =====================================================
station_summary = total_calls_by_station.merge(
    missing_by_station,
    on="Station",
    how="left"
)

station_summary = station_summary.merge(
    first_due_calls_by_station,
    on="Station",
    how="left"
)

station_summary = station_summary.merge(
    response_time_by_station_clean,
    on="Station",
    how="left"
)

station_summary = station_summary.merge(
    percent_over_480_by_station,
    on="Station",
    how="left"
)

for col in [
    "Insufficient Data Calls",
    "1st Due Calls",
    "1st Due Emergent Calls",
    "1st Due Emergent Calls >480"
]:
    if col in station_summary.columns:
        station_summary[col] = station_summary[col].fillna(0).astype(int)

station_summary = station_summary[
    [
        "Station",
        "Total Calls",
        "Insufficient Data Calls",
        "1st Due Calls",
        "1st Due Emergent Calls",
        "Average 1st Due Emergent Response Time",
        "90th Percentile 1st Due Emergent Response Time",
        "1st Due Emergent Calls >480",
        "Percent >480"
    ]
]

station_summary.columns = [
    "Station",
    "Total Calls",
    "Insufficient Data",
    "1st Due Calls",
    "Emergent 1st Due Calls",
    "Avg Response",
    "90th %",
    ">480 Calls",
    "% >480"
]

# =====================================================
# CONSOLE PRINTOUTS
# =====================================================
print("\n==============================")
print("DATA QUALITY")
print("==============================")
print(f"Call Log Raw Records: {raw_record_count}")
print(f"Call Log Duplicate Incident Numbers Removed: {duplicate_count}")
print(f"Call Log Records After Removing Duplicates: {deduped_record_count}")

print(f"Monthly Incident List Raw Records: {monthly_raw_record_count}")
print(f"Monthly Incident List Duplicate Incident Numbers Removed: {monthly_duplicate_count}")
print(f"Monthly Incident List Records After Removing Duplicates: {monthly_deduped_record_count}")

print(f"Calls in Monthly List Missing From Call Log: {missing_data_call_count}")

if missing_data_call_count > 0:
    print("\nMONTHLY CALLS MISSING FROM CALL LOG")
    print(missing_calls[[incident_col, station_col, "Data Status"]].to_string(index=False))

    print("\nINSUFFICIENT DATA CALLS BY STATION")
    print(missing_by_station.to_string(index=False))

print("\n==============================")
print("MONTHLY CALL COUNTS")
print("==============================")
print(f"Total Monthly Calls: {total_monthly_calls}")
print(f"Total Calls With Usable Call Log Records: {total_call_log_calls}")
print(f"Total Calls With Insufficient Data: {missing_data_call_count}")
print(f"Total 1st Due Calls: {total_first_due_calls}")
print(f"Total 1st Due Emergent Calls: {total_first_due_emergent_calls}")

print("\nTOTAL CALLS BY STATION")
print(total_calls_by_station.to_string(index=False))

print("\n1ST DUE CALLS BY STATION")
print(first_due_calls_by_station.to_string(index=False))

print("\n1ST DUE EMERGENT CALLS BY STATION")
print(first_due_emergent_calls_by_station.to_string(index=False))

print("\n==============================")
print("RESPONSE TIME - 1ST DUE EMERGENT ONLY")
print("==============================")
print(f"Average 1st Due Emergent Response Time: {avg_first_due_response_time}")
print(f"90th Percentile 1st Due Emergent Response Time: {p90_first_due_response_time}")

print("\n1ST DUE EMERGENT RESPONSE TIME BY STATION")
print(response_time_by_station_clean.to_string(index=False))

print("\n==============================")
print("MUTUAL AID")
print("==============================")
print(f"Total Mutual Aid Calls Given: {total_mutual_aid_calls}")

print("\nMUTUAL AID CALLS BY STATION")
print(mutual_aid_calls_by_station.to_string(index=False))

print("\n==============================")
print("1ST DUE EMERGENT RESPONSE TIME >480")
print("==============================")
print(f"Count of 1st Due Emergent Calls >480 Seconds: {count_first_due_over_480}")
print(f"Percent of 1st Due Emergent Calls >480 Seconds: {percent_first_due_over_480:.2f}%")

print("\n1ST DUE EMERGENT CALLS >480 BY STATION")
print(percent_over_480_by_station.to_string(index=False))

print("\n1ST DUE EMERGENT CALLS >480 DETAIL")
print(first_due_over_480_detail.to_string(index=False))

print("\n==============================")
print("1ST DUE EMERGENT >480 OVERLAP ANALYSIS")
print("==============================")
print(f"1st Due Emergent >480 Calls With Any Overlap: {first_due_over_480_with_any_overlap}")
print(f"Percent of 1st Due Emergent >480 Calls With Any Overlap: {percent_first_due_over_480_with_overlap:.2f}%")
print(f"Average Overlapping Calls Among 1st Due Emergent >480 Calls: {average_overlap_count_for_first_due_over_480:.2f}")

print("\n1ST DUE EMERGENT >480 OVERLAP SUMMARY BY STATION")
print(first_due_over_480_overlap_by_station_clean.to_string(index=False))

print("\n==============================")
print("SYSTEM STRESS HOURS")
print("==============================")
print(stress_hours_df.to_string(index=False))

print("\n==============================")
print("CALLS BY HOUR")
print("==============================")
print(calls_by_hour_df.to_string(index=False))

print(f"\nChart saved to: {chart_path}")

# =====================================================
# CREATE WORD DOCUMENT BRIEFING
# =====================================================
doc = Document()

# -----------------------------------------------------
# TITLE
# -----------------------------------------------------
title = doc.add_heading("District 7 Fire/Rescue", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading("Monthly Analysis Report", 2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# -----------------------------------------------------
# EXECUTIVE SUMMARY
# -----------------------------------------------------
doc.add_heading("Executive Summary", level=1)

doc.add_paragraph(f"Total Monthly Calls: {total_monthly_calls}")
doc.add_paragraph(f"Total Calls With Usable Call Log Records: {total_call_log_calls}")
doc.add_paragraph(f"Total 1st Due Calls: {total_first_due_calls}")
doc.add_paragraph(f"Total 1st Due Emergent Calls: {total_first_due_emergent_calls}")

doc.add_paragraph(
    f"{missing_data_call_count} calls were identified in the monthly incident list "
    f"but had insufficient call-log data for timing or performance analysis. "
    f"These calls were included in station call totals but excluded from response-time, "
    f"overlap, duration, and other performance metrics."
)

doc.add_paragraph(
    f"Average 1st Due Emergent Response Time: {avg_first_due_response_time}"
)

doc.add_paragraph(
    f"90th Percentile 1st Due Emergent Response Time: {p90_first_due_response_time}"
)

doc.add_paragraph(
    f"{count_first_due_over_480} emergent 1st due calls exceeded 480 seconds "
    f"({percent_first_due_over_480:.2f}%)."
)

doc.add_paragraph(
    f"These delayed responses occurred with an average of "
    f"{average_overlap_count_for_first_due_over_480:.2f} simultaneous incidents, "
    f"indicating increased system demand during performance degradation."
)

doc.add_paragraph(
    f"{percent_first_due_over_480_with_overlap:.2f}% of delayed emergent calls "
    f"occurred while at least one other incident was active."
)

# -----------------------------------------------------
# STATION PERFORMANCE TABLE
# -----------------------------------------------------
add_table_from_df(
    doc,
    station_summary,
    "Station Performance"
)

# -----------------------------------------------------
# MUTUAL AID TABLE
# -----------------------------------------------------
doc.add_heading("Mutual Aid Activity", level=1)

doc.add_paragraph(f"Total Mutual Aid Calls Given: {total_mutual_aid_calls}")

add_table_from_df(
    doc,
    mutual_aid_calls_by_station,
    title=None
)

# -----------------------------------------------------
# SYSTEM LOAD / OVERLAP TABLE
# -----------------------------------------------------
doc.add_heading("System Load & Overlapping Incidents", level=1)

doc.add_paragraph(
    f"Delayed emergent responses (>480 sec) were associated with an average of "
    f"{average_overlap_count_for_first_due_over_480:.2f} concurrent incidents. "
    f"Periods of system stress were identified where overlapping calls were greater than or equal to 2."
)

add_table_from_df(
    doc,
    first_due_over_480_overlap_by_station_clean,
    title=None
)

# -----------------------------------------------------
# CALL VOLUME DISTRIBUTION - CHART ONLY
# -----------------------------------------------------
doc.add_heading("Call Volume by Hour of Day", level=1)

doc.add_paragraph(
    "The following chart illustrates call distribution across the 24-hour period. "
    "Shaded areas denote hours where overlapping calls were greater than or equal to 2. "
    "Calls with insufficient timing data were excluded from this hourly chart."
)

doc.add_picture(chart_path, width=Inches(6.5))

# -----------------------------------------------------
# FOOTER
# -----------------------------------------------------
section = doc.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]

run_date = datetime.now().strftime("%B %d, %Y")

footer_para.text = f"Analysis by Michael Boulding | {run_date}"
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# -----------------------------------------------------
# SAVE DOCUMENT
# -----------------------------------------------------
doc.save(output_doc_path)

print(f"\nWord report saved to: {output_doc_path}")